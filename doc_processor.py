"""
doc_processor.py — Extraction de texte depuis PDF, DOCX et Images.

Utilise pypdf (PDF natif), python-docx (DOCX), et MarkItDown comme
extracteur universel de fallback. Pour les images, utilise un LLM
vision si une clé API est disponible.
"""

import os
import io
import re
from pathlib import Path
from typing import Optional

# ---------------------------------------------------------------------------
# Helpers internes
# ---------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    """Nettoie le texte extrait : supprime whitespace excessif."""
    if not text:
        return ""
    # Normaliser les sauts de ligne multiples
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Supprimer les espaces en fin de ligne
    text = "\n".join(line.rstrip() for line in text.splitlines())
    return text.strip()


def _detect_language(text: str) -> str:
    """Détecte la langue du texte (retourne 'fr' par défaut)."""
    try:
        from langdetect import detect
        if len(text) > 50:
            return detect(text[:500])
    except Exception:
        pass
    return "fr"


# ---------------------------------------------------------------------------
# Extracteurs par type
# ---------------------------------------------------------------------------

def extract_from_pdf(file_path: str) -> dict:
    """
    Extrait le texte d'un PDF via pypdf.
    Fallback sur MarkItDown si pypdf échoue ou retourne peu de texte.

    Returns:
        dict avec 'text', 'nb_pages', 'langue', 'methode'
    """
    text = ""
    nb_pages = 0
    method = "pypdf"

    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        nb_pages = len(reader.pages)
        parts = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            parts.append(page_text)
        text = "\n\n".join(parts)
    except Exception as e:
        method = "markitdown"

    # Fallback MarkItDown si le texte est insuffisant
    if len(text.strip()) < 100:
        try:
            from markitdown import MarkItDown
            md = MarkItDown()
            result = md.convert(file_path)
            text = result.text_content or text
            method = "markitdown"
        except Exception:
            pass

    text = _clean_text(text)
    return {
        "text": text,
        "nb_pages": nb_pages,
        "langue": _detect_language(text),
        "methode": method,
        "nb_chars": len(text),
    }


def extract_from_docx(file_path: str) -> dict:
    """
    Extrait le texte d'un fichier Word (.docx).

    Returns:
        dict avec 'text', 'nb_pages', 'langue', 'methode'
    """
    text = ""
    method = "python-docx"

    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Inclure aussi les tableaux
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        text = "\n\n".join(paragraphs)
    except Exception:
        # Fallback MarkItDown
        try:
            from markitdown import MarkItDown
            md = MarkItDown()
            result = md.convert(file_path)
            text = result.text_content or ""
            method = "markitdown"
        except Exception:
            pass

    text = _clean_text(text)
    return {
        "text": text,
        "nb_pages": 0,  # DOCX n'a pas de notion de pages
        "langue": _detect_language(text),
        "methode": method,
        "nb_chars": len(text),
    }


def extract_from_image(file_path: str, llm_client=None) -> dict:
    """
    Extrait le texte d'une image.
    
    Stratégie :
    1. Si un client LLM vision est disponible → description via API
    2. Sinon → MarkItDown (OCR basique)
    3. Sinon → message informatif

    Returns:
        dict avec 'text', 'nb_pages', 'langue', 'methode'
    """
    text = ""
    method = "unavailable"

    # Essai 1 : LLM vision
    if llm_client:
        try:
            import base64
            with open(file_path, "rb") as f:
                image_data = base64.b64encode(f.read()).decode("utf-8")

            ext = Path(file_path).suffix.lower().lstrip(".")
            mime_map = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png",
                        "gif": "gif", "webp": "webp"}
            mime = mime_map.get(ext, "jpeg")

            response = llm_client.describe_image(image_data, mime)
            text = response or ""
            method = "llm-vision"
        except Exception:
            pass

    # Essai 2 : MarkItDown
    if not text:
        try:
            from markitdown import MarkItDown
            md = MarkItDown()
            result = md.convert(file_path)
            text = result.text_content or ""
            method = "markitdown"
        except Exception:
            pass

    # Fallback
    if not text:
        text = (
            f"[Image : {Path(file_path).name}]\n"
            "Extraction automatique non disponible. "
            "Configurez un LLM vision (Gemini/GPT-4o) dans les Paramètres."
        )
        method = "unavailable"

    text = _clean_text(text)
    return {
        "text": text,
        "nb_pages": 1,
        "langue": _detect_language(text),
        "methode": method,
        "nb_chars": len(text),
    }


# ---------------------------------------------------------------------------
# Point d'entrée principal
# ---------------------------------------------------------------------------

def process_document(file_path: str, file_type: str = None,
                     llm_client=None) -> dict:
    """
    Extrait le texte d'un document selon son type.

    Args:
        file_path: Chemin absolu vers le fichier
        file_type: Extension ('pdf', 'docx', 'jpg', ...) ou None pour détecter
        llm_client: Instance LLMHandler optionnelle (pour vision)

    Returns:
        dict avec 'text', 'nb_pages', 'langue', 'methode', 'nb_chars', 'ok'
    """
    path = Path(file_path)

    if file_type is None:
        file_type = path.suffix.lower().lstrip(".")

    result = {"text": "", "nb_pages": 0, "langue": "fr",
              "methode": "none", "nb_chars": 0, "ok": False}

    if not path.exists():
        result["error"] = f"Fichier introuvable : {file_path}"
        return result

    try:
        if file_type == "pdf":
            data = extract_from_pdf(str(path))
        elif file_type in ("docx", "doc"):
            data = extract_from_docx(str(path))
        elif file_type in ("jpg", "jpeg", "png", "gif", "bmp", "webp", "tiff"):
            data = extract_from_image(str(path), llm_client)
        else:
            # Tentative générique avec MarkItDown
            try:
                from markitdown import MarkItDown
                md = MarkItDown()
                res = md.convert(str(path))
                data = {
                    "text": _clean_text(res.text_content or ""),
                    "nb_pages": 0,
                    "langue": "fr",
                    "methode": "markitdown",
                    "nb_chars": len(res.text_content or ""),
                }
            except Exception:
                data = {
                    "text": f"[Format non supporté : {file_type}]",
                    "nb_pages": 0,
                    "langue": "fr",
                    "methode": "none",
                    "nb_chars": 0,
                }

        result.update(data)
        result["ok"] = len(result["text"]) > 0

    except Exception as e:
        result["error"] = str(e)

    return result


def get_file_size_kb(file_path: str) -> float:
    """Retourne la taille d'un fichier en kilo-octets."""
    try:
        return round(Path(file_path).stat().st_size / 1024, 2)
    except Exception:
        return 0.0


SUPPORTED_EXTENSIONS = {
    "pdf": "📄 PDF",
    "docx": "📝 Word",
    "doc": "📝 Word (ancien)",
    "jpg": "🖼️ Image JPG",
    "jpeg": "🖼️ Image JPEG",
    "png": "🖼️ Image PNG",
    "gif": "🖼️ Image GIF",
    "webp": "🖼️ Image WebP",
    "txt": "📃 Texte",
    "md": "📃 Markdown",
}
