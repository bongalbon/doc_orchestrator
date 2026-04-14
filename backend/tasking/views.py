import io
import os
import urllib.request
import json
from django.http import FileResponse
from django.shortcuts import get_object_or_404
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from openpyxl import Workbook
from tasking.realtime import broadcast_activity
from rest_framework import mixins, status, viewsets
from rest_framework.decorators import action
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response

from agents.models import Agent
from tasking.models import (
    AgentTask,
    AuditLog,
    Workflow,
    Notification,
    WorkflowStep,
    ProviderCredential,
)
from tasking.permissions import IsOperatorOrReadOnly
from tasking.serializers import (
    AgentTaskSerializer,
    TaskCreateSerializer,
    WorkflowSerializer,
    NotificationSerializer,
    WorkflowCreateSerializer,
    ProviderCredentialSerializer,
)
from tasking.services import (
    active_agents_snapshot,
    enqueue_task,
    running_snapshot,
    start_workflow,
)
from tasking.tasks import execute_agent_task


class AgentTaskViewSet(
    mixins.CreateModelMixin,
    mixins.ListModelMixin,
    mixins.RetrieveModelMixin,
    mixins.UpdateModelMixin,
    viewsets.GenericViewSet,
):
    serializer_class = AgentTaskSerializer
    queryset = AgentTask.objects.select_related(
        "requested_agent", "assigned_agent"
    ).all()
    permission_classes = [IsAuthenticated, IsOperatorOrReadOnly]

    def get_queryset(self):
        qs = super().get_queryset()
        status = self.request.query_params.get("status")
        if status:
            qs = qs.filter(status=status)
        return qs

    def create(self, request, *args, **kwargs):
        create_serializer = TaskCreateSerializer(data=request.data)
        create_serializer.is_valid(raise_exception=True)
        payload = create_serializer.validated_data

        requested_agent = None
        requested_agent_id = payload.get("requested_agent_id")
        if requested_agent_id is not None:
            requested_agent = get_object_or_404(
                Agent, id=requested_agent_id, is_active=True
            )

        task = enqueue_task(
            title=payload["title"],
            prompt=payload["prompt"],
            requested_agent=requested_agent,
            timeout_seconds=payload["timeout_seconds"],
            actor=request.user,
            provider=payload["provider"],
            model_name=payload["model_name"],
            api_key=payload.get("api_key", ""),
            ollama_url=payload.get("ollama_url", ""),
        )
        serialized = AgentTaskSerializer(task)
        return Response(serialized.data, status=status.HTTP_201_CREATED)

    @action(detail=False, methods=["get"], url_path="activity")
    def activity(self, request):
        return Response(
            {
                "running_tasks": running_snapshot(),
                "active_agents": active_agents_snapshot(),
            }
        )

    @action(detail=True, methods=["post"], url_path="cancel")
    def cancel(self, request, pk=None):
        task = self.get_object()
        if task.status in {"done", "failed", "cancelled"}:
            return Response(
                {"detail": "Task already finished."}, status=status.HTTP_400_BAD_REQUEST
            )

        task.cancel_requested = True
        task.mark_cancelled()  # Sets status to 'cancelled' and saves

        if task.celery_task_id:
            try:
                execute_agent_task.AsyncResult(task.celery_task_id).revoke(
                    terminate=True, signal="SIGKILL"
                )
            except Exception:
                pass

        AuditLog.objects.create(
            action="task_cancel_requested", actor=request.user, task=task
        )
        broadcast_activity({"event": "task_cancelled", "task_id": task.id})

        return Response({"ok": True})

    @action(detail=True, methods=["post"], url_path="retry")
    def retry(self, request, pk=None):
        task = self.get_object()
        if task.status not in {"failed", "cancelled"}:
            return Response(
                {"detail": "Only failed/cancelled tasks can be retried."}, status=400
            )
        if task.retry_count >= task.max_retries:
            return Response({"detail": "Retry limit reached."}, status=400)
        task.retry_count += 1
        task.status = "queued"
        task.cancel_requested = False
        task.error_message = ""
        task.save(
            update_fields=[
                "retry_count",
                "status",
                "cancel_requested",
                "error_message",
                "updated_at",
            ]
        )
        job = execute_agent_task.apply_async(
            args=[task.id],
            time_limit=task.timeout_seconds,
            soft_time_limit=max(10, task.timeout_seconds - 5),
        )
        task.celery_task_id = job.id or ""
        task.save(update_fields=["celery_task_id", "updated_at"])
        AuditLog.objects.create(
            action="task_retried",
            actor=request.user,
            task=task,
            metadata={"job_id": job.id},
        )
        return Response({"ok": True})

    @action(detail=False, methods=["get"], url_path="ollama-models")
    def ollama_models(self, request):
        custom_url = request.query_params.get("url")
        ollama_url = custom_url or os.getenv(
            "OLLAMA_BASE_URL", "http://host.docker.internal:11434"
        )
        try:
            with urllib.request.urlopen(
                f"{ollama_url}/api/tags", timeout=4
            ) as response:
                if response.status == 200:
                    data = json.loads(response.read().decode())
                    models = [m["name"] for m in data.get("models", [])]
                    return Response({"models": models})
        except Exception as e:
            return Response(
                {"models": [], "error": str(e)},
                status=status.HTTP_503_SERVICE_UNAVAILABLE,
            )
        return Response({"models": []})

    @action(detail=False, methods=["get"], url_path="provider-models")
    def provider_models(self, request):
        provider = request.query_params.get("provider", "gemini")
        # En fonction du provider, on peut renvoyer une liste hardcodée augmentée
        # ou appeler LiteLLM/API directes si on a une clé active pour l'user.

        # Pour Gemini, l'user veut spécifiquement ces versions (même si futures/expérimentales)
        if provider == "gemini":
            return Response(
                {
                    "models": [
                        "gemini-3.1-pro",
                        "gemini-3.0-flash",
                        "gemini-2.5-pro",
                        "gemini-2.5-flash",
                        "gemini-2.5-flash-lite",
                        "gemini-2.0-flash",
                        "gemini-2.0-flash-lite-preview-02-05",
                        "gemini-1.5-pro",
                        "gemini-1.5-flash",
                        "gemini-1.5-flash-8b",
                    ]
                }
            )

        # Pour les autres, on garde une base solide en attendant une intégration plus poussée
        defaults = {
            "openai": ["o1-mini", "o1-preview", "gpt-4o", "gpt-4o-mini", "gpt-4-turbo"],
            "anthropic": [
                "claude-3-7-sonnet-20250219",
                "claude-3-5-sonnet-20241022",
                "claude-3-opus-20240229",
            ],
            "grok": ["grok-2-1212", "grok-beta"],
        }
        return Response({"models": defaults.get(provider, [])})

    @action(detail=True, methods=["get"], url_path="export")
    def export_document(self, request, pk=None):
        task = self.get_object()
        fmt = request.query_params.get("fmt", "docx")
        content = task.result or "Aucun contenu généré."
        buffer = io.BytesIO()

        if fmt == "docx":
            doc = Document()
            doc.add_heading(task.title, 0)
            doc.add_paragraph(content)
            doc.save(buffer)
            filename = f"export_{pk}.docx"
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif fmt == "pdf":
            p = canvas.Canvas(buffer, pagesize=A4)
            width, height = A4
            p.setFont("Helvetica", 12)
            p.drawString(50, height - 50, task.title)
            textobject = p.beginText(50, height - 70)
            for line in content.split("\n"):
                if len(line) > 90:
                    for i in range(0, len(line), 90):
                        textobject.textLine(line[i : i + 90])
                else:
                    textobject.textLine(line)
            p.drawText(textobject)
            p.showPage()
            p.save()
            filename = f"export_{pk}.pdf"
            content_type = "application/pdf"
        elif fmt == "xlsx":
            wb = Workbook()
            ws = wb.active
            ws.title = "Resultat"
            ws.append(["Titre", task.title])
            ws.append(
                [
                    "Date",
                    task.finished_at.strftime("%Y-%m-%d") if task.finished_at else "",
                ]
            )
            ws.append([])
            ws.append(["Contenu:"])
            for line in content.split("\n"):
                ws.append([line])
            wb.save(buffer)
            filename = f"export_{pk}.xlsx"
            content_type = (
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        else:
            return Response(
                {"detail": "Format non supporté."}, status=status.HTTP_400_BAD_REQUEST
            )

        buffer.seek(0)
        return FileResponse(
            buffer, as_attachment=True, filename=filename, content_type=content_type
        )


class WorkflowViewSet(viewsets.ModelViewSet):
    queryset = (
        Workflow.objects.prefetch_related("steps")
        .select_related("manager_agent", "user")
        .all()
    )
    serializer_class = WorkflowSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = super().get_queryset()
        status = self.request.query_params.get("status")
        if status:
            qs = qs.filter(status=status)
        return qs

    def create(self, request, *args, **kwargs):
        ser = WorkflowCreateSerializer(data=request.data)
        ser.is_valid(raise_exception=True)
        payload = ser.validated_data

        manager_agent = None
        if payload.get("manager_agent_id"):
            manager_agent = get_object_or_404(Agent, id=payload["manager_agent_id"])

        workflow = start_workflow(
            title=payload["title"],
            prompt=payload["prompt"],
            manager_agent_id=payload.get("manager_agent_id"),
            user=request.user,
            provider=payload.get("provider", "gemini"),
            model_name=payload.get("model_name", "gemini-2.0-flash"),
            ollama_url=payload.get("ollama_url", ""),
        )
        return Response(
            WorkflowSerializer(workflow).data, status=status.HTTP_201_CREATED
        )

    @action(detail=True, methods=["post"], url_path="approve")
    def approve(self, request, pk=None):
        workflow = self.get_object()
        notification = workflow.notifications.filter(status="pending").first()
        if notification:
            notification.status = "approved"
            notification.save()
        workflow.status = "completed"
        workflow.save()
        return Response({"status": "completed"})

    @action(detail=True, methods=["post"], url_path="cancel")
    def cancel(self, request, pk=None):
        workflow = self.get_object()
        if workflow.status in {"completed", "failed"}:
            return Response({"detail": "Workflow already finished."}, status=400)

        workflow.cancel_requested = True
        workflow.status = "failed"
        workflow.final_result = "Annulé par l'utilisateur."
        workflow.save(
            update_fields=["cancel_requested", "status", "final_result", "updated_at"]
        )

        if workflow.celery_task_id:
            try:
                from tasking.tasks import run_workflow_orchestration

                run_workflow_orchestration.AsyncResult(workflow.celery_task_id).revoke(
                    terminate=True, signal="SIGKILL"
                )
            except Exception:
                pass

        AuditLog.objects.create(
            action="workflow_cancelled",
            actor=request.user,
            metadata={"workflow_id": workflow.id},
        )
        broadcast_activity({"event": "workflow_cancelled", "workflow_id": workflow.id})

        return Response({"ok": True})

    @action(detail=True, methods=["post"], url_path="reject")
    def reject(self, request, pk=None):
        workflow = self.get_object()
        feedback = request.data.get("feedback", "No feedback provided.")
        notification = workflow.notifications.filter(status="pending").first()
        if notification:
            notification.status = "rejected"
            notification.user_feedback = feedback
            notification.save()

        # Reset workflow to thinking to re-run iteration
        workflow.status = "thinking"
        workflow.save()

        from tasking.tasks import run_workflow_orchestration

        run_workflow_orchestration.apply_async(args=[workflow.id])
        return Response({"status": "resubmitted"})


class NotificationViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Notification.objects.all()
    serializer_class = NotificationSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return self.queryset.filter(user=self.request.user)


class CredentialViewSet(viewsets.ModelViewSet):
    serializer_class = ProviderCredentialSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return ProviderCredential.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)
