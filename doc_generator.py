"""
doc_generator.py — Génération de documents depuis du Markdown.

Convertit le Markdown en PDF (via ReportLab) et en DOCX (via python-docx).
Pas de dépendance à LaTeX ou Pandoc — 100% compatible Termux.
"""

import re
import io
from pathlib import Path
from datetime import datetime

from config import EXPORTS_DIR


# ---------------------------------------------------------------------------
# Helpers Markdown → structure
# ---------------------------------------------------------------------------

def _parse_markdown_lines(md_text: str) -> list[dict]:
    """
    Parse le Markdown ligne par ligne et retourne une liste de tokens structurés.
    Types reconnus : h1, h2, h3, paragraph, bullet, numbered, hr, code, bold_line, empty
    """
    tokens = []
    lines = md_text.splitlines()
    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.rstrip()

        # Blocs de code
        if stripped.startswith("```"):
            if in_code_block:
                tokens.append({"type": "code", "text": "\n".join(code_lines)})
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(stripped)
            continue

        # Titres
        if stripped.startswith("### "):
            tokens.append({"type": "h3", "text": stripped[4:].strip()})
        elif stripped.startswith("## "):
            tokens.append({"type": "h2", "text": stripped[3:].strip()})
        elif stripped.startswith("# "):
            tokens.append({"type": "h1", "text": stripped[2:].strip()})
        # Séparateur horizontal
        elif stripped in ("---", "___", "***"):
            tokens.append({"type": "hr"})
        # Listes à puces
        elif stripped.startswith("- ") or stripped.startswith("* "):
            tokens.append({"type": "bullet", "text": stripped[2:].strip()})
        # Listes numérotées
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped).strip()
            tokens.append({"type": "numbered", "text": text})
        # Ligne vide
        elif not stripped:
            tokens.append({"type": "empty"})
        # Paragraphe standard
        else:
            tokens.append({"type": "paragraph", "text": stripped})

    return tokens


def _strip_md_inline(text: str) -> str:
    """Supprime le formatage inline Markdown (gras, italique, liens)."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


# ---------------------------------------------------------------------------
# Génération PDF via ReportLab
# ---------------------------------------------------------------------------

def generate_pdf(md_text: str, title: str = "Document") -> bytes:
    """
    Convertit du Markdown en PDF via ReportLab.

    Args:
        md_text: Contenu Markdown
        title: Titre du document (métadonnée PDF)

    Returns:
        bytes du fichier PDF
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        HRFlowable, ListFlowable, ListItem, Preformatted
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=title,
        author="AI-Doc-Orchestrator",
    )

    # Styles
    styles = getSampleStyleSheet()

    style_h1 = ParagraphStyle(
        "H1",
        fontSize=20,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#1a237e"),
        spaceAfter=14,
        spaceBefore=20,
        alignment=TA_CENTER,
    )
    style_h2 = ParagraphStyle(
        "H2",
        fontSize=15,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#283593"),
        spaceAfter=10,
        spaceBefore=16,
        borderPad=4,
    )
    style_h3 = ParagraphStyle(
        "H3",
        fontSize=12,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#3949ab"),
        spaceAfter=8,
        spaceBefore=12,
    )
    style_body = ParagraphStyle(
        "Body",
        fontSize=10,
        fontName="Helvetica",
        leading=16,
        spaceAfter=8,
        alignment=TA_JUSTIFY,
    )
    style_bullet = ParagraphStyle(
        "Bullet",
        fontSize=10,
        fontName="Helvetica",
        leading=14,
        leftIndent=20,
        spaceAfter=4,
        bulletIndent=10,
    )
    style_code = ParagraphStyle(
        "Code",
        fontSize=9,
        fontName="Courier",
        backColor=colors.HexColor("#f5f5f5"),
        leftIndent=15,
        rightIndent=15,
        spaceAfter=10,
        leading=13,
    )

    # Construire les éléments
    story = []
    tokens = _parse_markdown_lines(md_text)
    bullet_buffer = []
    numbered_buffer = []

    def flush_lists():
        nonlocal bullet_buffer, numbered_buffer
        if bullet_buffer:
            items = [
                ListItem(
                    Paragraph(_strip_md_inline(t), style_bullet),
                    bulletColor=colors.HexColor("#3949ab"),
                    leftIndent=20,
                )
                for t in bullet_buffer
            ]
            story.append(ListFlowable(items, bulletType="bullet", start="•"))
            story.append(Spacer(1, 6))
            bullet_buffer = []
        if numbered_buffer:
            items = [
                ListItem(
                    Paragraph(_strip_md_inline(t), style_bullet),
                    bulletColor=colors.HexColor("#3949ab"),
                    leftIndent=20,
                )
                for t in numbered_buffer
            ]
            story.append(ListFlowable(items, bulletType="1", start=1))
            story.append(Spacer(1, 6))
            numbered_buffer = []

    for token in tokens:
        ttype = token["type"]
        text = token.get("text", "")
        clean = _strip_md_inline(text)

        if ttype == "bullet":
            numbered_buffer and flush_lists()
            bullet_buffer.append(text)
            continue
        elif ttype == "numbered":
            bullet_buffer and flush_lists()
            numbered_buffer.append(text)
            continue
        else:
            flush_lists()

        if ttype == "h1":
            story.append(Paragraph(clean, style_h1))
        elif ttype == "h2":
            story.append(HRFlowable(width="100%", thickness=1,
                                    color=colors.HexColor("#3949ab"),
                                    spaceAfter=4))
            story.append(Paragraph(clean, style_h2))
        elif ttype == "h3":
            story.append(Paragraph(clean, style_h3))
        elif ttype == "hr":
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=colors.HexColor("#cccccc"),
                                    spaceBefore=8, spaceAfter=8))
        elif ttype == "code":
            story.append(Preformatted(text, style_code))
        elif ttype == "paragraph" and clean:
            story.append(Paragraph(clean, style_body))
        elif ttype == "empty":
            story.append(Spacer(1, 6))

    flush_lists()

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


# ---------------------------------------------------------------------------
# Génération DOCX via python-docx
# ---------------------------------------------------------------------------

def generate_docx(md_text: str, title: str = "Document") -> bytes:
    """
    Convertit du Markdown en DOCX via python-docx.

    Args:
        md_text: Contenu Markdown
        title: Titre du document

    Returns:
        bytes du fichier DOCX
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Propriétés du document
    doc.core_properties.title = title
    doc.core_properties.author = "AI-Doc-Orchestrator"

    tokens = _parse_markdown_lines(md_text)

    def add_styled_paragraph(d, text, style_name, bold=False,
                             font_size=11, color=None, alignment=None,
                             space_before=6, space_after=6):
        """Ajoute un paragraphe formaté."""
        p = d.add_paragraph()
        p.style = d.styles["Normal"]
        run = p.add_run(_strip_md_inline(text))
        run.bold = bold
        run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        if alignment:
            p.alignment = alignment
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    for token in tokens:
        ttype = token["type"]
        text = token.get("text", "")

        if ttype == "h1":
            add_styled_paragraph(
                doc, text, "Heading 1",
                bold=True, font_size=18,
                color=(26, 35, 126),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=12, space_after=10,
            )
        elif ttype == "h2":
            p = add_styled_paragraph(
                doc, text, "Heading 2",
                bold=True, font_size=14,
                color=(40, 53, 147),
                space_before=14, space_after=8,
            )
            p.paragraph_format.keep_with_next = True
        elif ttype == "h3":
            add_styled_paragraph(
                doc, text, "Heading 3",
                bold=True, font_size=12,
                color=(57, 73, 171),
                space_before=10, space_after=6,
            )
        elif ttype == "bullet":
            p = doc.add_paragraph(
                _strip_md_inline(text),
                style="List Bullet"
            )
            p.paragraph_format.space_after = Pt(3)
        elif ttype == "numbered":
            p = doc.add_paragraph(
                _strip_md_inline(text),
                style="List Number"
            )
            p.paragraph_format.space_after = Pt(3)
        elif ttype == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.border_bottom = True
        elif ttype == "code":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)
        elif ttype == "paragraph" and text.strip():
            p = doc.add_paragraph()
            # Gestion du gras inline **texte**
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(re.sub(r"\*(.+?)\*", r"\1", part))
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif ttype == "empty":
            doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ---------------------------------------------------------------------------
# Fonctions d'export vers fichier
# ---------------------------------------------------------------------------

def save_pdf(md_text: str, filename: str) -> Path:
    """Génère et sauvegarde un PDF dans le dossier exports."""
    pdf_bytes = generate_pdf(md_text, title=filename)
    output_path = EXPORTS_DIR / f"{filename}.pdf"
    output_path.write_bytes(pdf_bytes)
    return output_path


def save_docx(md_text: str, filename: str) -> Path:
    """Génère et sauvegarde un DOCX dans le dossier exports."""
    docx_bytes = generate_docx(md_text, title=filename)
    output_path = EXPORTS_DIR / f"{filename}.docx"
    output_path.write_bytes(docx_bytes)
    return output_path


def generate_filename(titre: str, type_doc: str) -> str:
    """Génère un nom de fichier propre depuis le titre et le type de doc."""
    date_str = datetime.now().strftime("%Y%m%d_%H%M")
    clean_title = re.sub(r"[^\w\s-]", "", titre)
    clean_title = re.sub(r"\s+", "_", clean_title.strip())[:40]
    clean_type = re.sub(r"\s+", "_", type_doc)[:20]
    return f"{clean_type}_{clean_title}_{date_str}"
